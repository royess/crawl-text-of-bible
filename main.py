import requests
import re
from docx import Document


def write_in_all(urls, doc):

    for i in range(len(urls)):
        # write title(ch)
        print("Start to write in ch" + str(i+1))
        p = doc.add_paragraph('')
        p.add_run(str(i+1))

        html = requests.get(urls[i])
        text = html.text
        terms = re.findall("<p><a href=\".*?\" title='.*?'><span id='\w' class='versehover'>(\w*?)"
                           " </span></a>(.+?)</p>", text)

        for term in terms:
            p = doc.add_paragraph('')
            p.add_run(term[0])

            p = doc.add_paragraph('')
            words = term[1].split()
            for word in words:
                r = re.search("<em>(.\w*)", word)
                r1 = re.search("(.\w*)</em>", word)
                if r:
                    print(r.group(1))
                    run = p.add_run(r.group(1) + " ")
                    run.italic = True
                elif r1:
                    print(r1.group(1))
                    run = p.add_run(r1.group(1) + " ")
                    run.italic = True
                else:
                    print(word)
                    p.add_run(word + " ")



if __name__ == "__main__":

    curls = ["https://www.kingjamesbibleonline.org/1-Corinthians-Chapter-" + str(i) + "/"
             for i in range(1, 17)]

    jurls = ["https://www.kingjamesbibleonline.org/1-John-Chapter-" + str(i) + "/"
             for i in range(1, 6)]

    doc = Document()

    write_in_all(curls, doc)
    write_in_all(jurls, doc)

    doc.save('res.docx')







